import os
from datetime import datetime
import docx
from docx.shared import Pt
from subprocess import call
from sys import exec_prefix, executable


class DocParser:
    def __init__(self, pdf_file=""):
        self.datetime = datetime.now()
        self.date = f"{self.datetime.day}.{self.datetime.month}.{self.datetime.year}"
        self.time = f"{self.datetime.hour}:{self.datetime.minute}:{self.datetime.second}"

        [os.makedirs(folder, exist_ok=True) for folder in ["docx", "logs", "txt"]]

        self.filename = pdf_file.strip(".pdf")
        self.pdf_filename = pdf_file
        self.txt_filename = f"{self.filename}.txt"
        self.docx_filename = f"{self.filename}.docx"

        """Store data from pdf file."""
        self.txt_string = self.oznaczenie = self.krs = ""
        self.woj = self.powiat = self.gmina = self.miejsc = ""
        self.nazwa = self.regon = self.nip = ""
        self.nazwa_parsed = self.reg_nip_parsed = False

        """Use 'self.logfile.add_paragraph' method to add to a logfile."""
        self.logfile = docx.Document()

        """All variables will be parsed through the template document."""
        self.template_path = "template.docx"
        self.document = docx.Document(self.template_path)
        self.font = self.document.styles['Normal'].font
        self.font.name = 'Arial'

        """Locks the cell in the document after parsing."""
        self.oznaczenie_done = self.woj_done = self.powiat_done = self.gmina_done = self.miejsc_done = False
        self.krs_done = self.nazwa_done = self.regon_done = self.nip_done = self.data_done = False

        self.regon_full = self.nip_full = ""  # bugfix

    def extract_pdf(self, env="default"):
        if env is "default":
            call([executable,
                  os.path.join(f"{exec_prefix}", "Scripts", "pdf2txt.py"),
                  os.path.join("pdf", f"{self.pdf_filename}"),
                  os.path.join(f"-otxt", f"{self.txt_filename}")])
        if env is "venv":
            call([os.path.join("venv", "Scripts", "python.exe"),
                  os.path.join("venv", "Scripts", "pdf2txt.py"),
                  os.path.join("pdf", f"{self.pdf_filename}"),
                  os.path.join(f"-otxt", f"{self.txt_filename}")])
        print(f"\n>>> Extracting text from '{self.pdf_filename}'")

    def open_txt(self):  # Initiates the string object from the txt file.
        with open(os.path.join("txt", f"{self.txt_filename}"), "r", encoding="utf-8") as file:
            self.txt_string = [line.rstrip('\n') for line in file]

    def debug(self, span=(0, 0)):  # Prints a line with index tag, spanning from [0] to [1].
        for index, line in enumerate(self.txt_string):
            try:
                if span[0] < index < span[1]:
                    print(index, line)
            except TypeError:
                print("Span must be a (x,y) touple.")

    def parse_txt(self):
        """Loops through the txt file and maps distinct values to the class variables"""
        lines = self.txt_string
        for line in lines:
            if line == "Oznaczenie sądu":
                index = lines.index(line) + 4
                oznaczenie = lines[index]
                check = lines[index + 1]
                if check is not "":  # checks for a newline (bugfix)
                    oznaczenie += f" {check}"
                self.oznaczenie = oznaczenie
            if line.startswith("kraj "):
                index = lines.index(line)
                try:
                    kraj, woj, powiat, gmina, miejsc = lines[index].split(",")
                    self.woj = woj.strip("woj. ")
                    self.powiat = powiat.strip("  powiat ")
                    self.gmina = gmina.strip("  gmina ")
                    miejsc = miejsc.strip("  miejsc.")
                    self.miejsc = miejsc
                    if miejsc == "":  # checks for a newline
                        index = index + 1
                        self.miejsc = lines[index]
                except ValueError as error:  # this should be in tests :/
                    print(f">Unknown error:: {error}")
            if line.startswith("Numer KRS"):
                krs = line.split()[-1]
                if len(krs) is not 10:
                    print(">>>!!! INVALID KRS !!!")
                    krs = "ERROR"  # This cannot be invalid, user alarm.
                self.krs = krs
            if line.startswith("3.") and not self.nazwa_parsed:
                self.nazwa_parsed = True
                index = lines.index(line) + 2
                self.nazwa = lines[index]
            if line.startswith("2.") and not self.reg_nip_parsed:
                self.reg_nip_parsed = True
                index = lines.index(line) + 2
                pola = lines[index]
                regon, nip = pola.split(",")
                regon = regon.strip("REGON: ")
                self.regon = regon.split()[0]
                self.regon_full = regon
                nip = nip.strip("  NIP: ")
                self.nip = nip.split()[0].strip("---")
                self.nip_full = nip

    def print_formatted_data(self):
        print(
            f">>> Parsing... '{self.filename}'\n"
            f"> Nazwa sądu: {self.oznaczenie},\n> Województwo: {self.woj}, Powiat: {self.powiat},"
            f" Gmina: {self.gmina}, Miejscowość: {self.miejsc},"
            f"\n> Firma spółki: {self.nazwa},\n> Numer KRS: {self.krs}, REGON: {self.regon_full}, NIP: {self.nip_full}"
        )

    def create_logfile(self):
        self.logfile.add_paragraph(
            f"{self.docx_filename} LOGFILE: \n\n{self.time}\n{self.date}\n\n"
            f"Numer KRS: {self.krs}\nNazwa sądu: {self.oznaczenie}\nWojewództwo: {self.woj}\n"
            f"Powiat: {self.powiat}\nGmina: {self.gmina}\n"
            f"Miejscowość: {self.miejsc}\nFirma spółki: {self.nazwa}\n"
            f"REGON: {self.regon_full}\nNIP: {self.nip_full}"
        )

    def save_logfile(self):
        self.logfile.save(os.path.join("logs", f"({self.datetime.day}-{self.datetime.month}-{self.datetime.year})"
                          f"-({self.datetime.hour}-{self.datetime.minute}-{self.datetime.second})-{self.docx_filename}"))
        print(">>> Logfile created!")

    def parse_docx(self):
        """loops through the list of all cells in the document.
         Looks for a variable name and inserts the new value in a paragraph"""
        for table in range(len(self.document.tables)):
            try:
                for cell in self.document.tables[table]._cells:
                    if cell.paragraphs[0].text.startswith("1.") and not self.oznaczenie_done:
                        cell.paragraphs[1].style.font.size = Pt(10)
                        cell.paragraphs[1].text = self.oznaczenie
                        self.oznaczenie_done = True

                    if cell.paragraphs[0].text.startswith("2.") and not self.woj_done:
                        cell.paragraphs[1].style.font.size = Pt(10)
                        cell.paragraphs[1].text = self.woj
                        self.woj_done = True

                    if cell.paragraphs[0].text.startswith("3.") and not self.powiat_done:
                        cell.paragraphs[1].style.font.size = Pt(10)
                        cell.paragraphs[1].text = self.powiat
                        self.powiat_done = True

                    if cell.paragraphs[0].text.startswith("4.") and not self.gmina_done:
                        cell.paragraphs[1].style.font.size = Pt(10)
                        cell.paragraphs[1].text = self.gmina
                        self.gmina_done = True

                    if cell.paragraphs[0].text.startswith("5.") and not self.miejsc_done:
                        cell.paragraphs[1].style.font.size = Pt(10)
                        cell.paragraphs[1].text = self.miejsc
                        self.miejsc_done = True

                    if cell.paragraphs[0].text.startswith("<KRS>"):
                        cell.paragraphs[0].text = ""
                        if len(self.krs) is not 10:
                            print("> !!! INVALID KRS !!!")
                        for liczba in self.krs:
                            paragraph = cell.paragraphs[0].add_run(liczba)
                            paragraph.font.size = Pt(10)
                            if int(liczba) < 5:
                                for i in range(15):
                                    paragraph_space = cell.paragraphs[0].add_run(" ")
                                    paragraph_space.font.size = Pt(2)
                            else:
                                for i in range(16):
                                    paragraph_space = cell.paragraphs[0].add_run(" ")
                                    paragraph_space.font.size = Pt(2)

                    if cell.paragraphs[0].text.startswith("8.") and not self.nazwa_done:
                        cell.paragraphs[1].style.font.size = Pt(10)
                        cell.paragraphs[1].text = f"   {self.nazwa}"
                        self.nazwa_done = True

                    if cell.paragraphs[0].text.startswith("<NIP>") and not self.nip_done:
                        cell.paragraphs[0].text = ""
                        if len(self.nip) is not 10:
                            print("> !!! INVALID NIP !!!")
                            self.nip = "               "
                        for liczba in self.nip:
                            try:
                                paragraph = cell.paragraphs[0].add_run(liczba)
                                paragraph.font.size = Pt(10)
                                for i in range(15):
                                    paragraph_space = cell.paragraphs[0].add_run(" ")
                                    paragraph_space.font.size = Pt(2)
                            except ValueError:
                                cell.paragraphs[0].add_run("")
                        self.nip_done = True

                    if cell.paragraphs[0].text.startswith("<REGON>") and not self.regon_done:
                        cell.paragraphs[0].text = ""
                        if len(self.regon) is not 9:
                            print("> !!! INVALID REGON !!!")
                        for liczba in self.regon:
                            try:
                                paragraph = cell.paragraphs[0].add_run(liczba)
                                paragraph.font.size = Pt(10)
                                if int(liczba) < 5:  # this block controls the number of whitespaces between numbers
                                    for i in range(15):
                                        paragraph_space = cell.paragraphs[0].add_run(" ")
                                        paragraph_space.font.size = Pt(2)
                                else:
                                    for i in range(16):
                                        paragraph_space = cell.paragraphs[0].add_run(" ")
                                        paragraph_space.font.size = Pt(2)
                            except ValueError:
                                cell.paragraphs[0].add_run("")
                        self.regon_done = True
                    if cell.paragraphs[0].text.startswith("DZ.MI"):
                        cell.paragraphs[0].text = self.date
            except IndexError:
                pass

    def save_docx(self):
        os.chdir(os.path.join("docx", ""))
        self.document.save(self.docx_filename)
        os.chdir(os.path.join("..", ""))
        print(f">>> File '{self.docx_filename}' created successfully! <\n")

    def clear_temp(self):  # removes .txt file
        os.remove(os.path.join("txt", f"{self.txt_filename}"))

    def parse_all(self):
        #  self.extract_pdf()  # extracts unicode text from pdf file
        self.open_txt()  # initiates temp text file
        self.parse_txt()  # parses temp text file
        self.print_formatted_data()  # console info
        self.parse_docx()  # inserts data into docx document
        self.create_logfile()  # create log object
        self.save_logfile()  # saves log object to a file
        self.save_docx()  # saves the .docx document


if __name__ == "__main__":  # debug
    parser = DocParser("odpis_aktualny_1.pdf")
    parser.extract_pdf()
    parser.open_txt()
    parser.clear_temp()
    parser.parse_txt()
    parser.print_formatted_data()
    parser.parse_docx()
    parser.save_docx()
    parser.create_logfile()
    parser.save_logfile()
