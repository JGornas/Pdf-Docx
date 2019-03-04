import docx
from docx.shared import Pt
import os
import datetime
import subprocess


class DocParser:
    def __init__(self, pdf_file=""):
        self.date = self.get_date()
        self.pdf_file = pdf_file
        self.filename = pdf_file.strip(".pdf")
        self.txt_file = f"{pdf_file}.txt"
        self.docx_file = f"{self.filename}.docx"

        self.txt_string = self.oznaczenie = self.krs = ""
        self.woj = self.powiat = self.gmina = self.miejsc = ""
        self.nazwa = self.regon = self.nip = ""
        self.nazwa_parsed = self.reg_nip_parsed = False

        self.template_path = "template.docx"
        self.document = docx.Document(self.template_path)  # creates docx document object
        self.font = self.document.styles['Normal'].font
        self.font.name = 'Arial'

        self.oznaczenie_done = self.woj_done = self.powiat_done = self.gmina_done = self.miejsc_done = False
        self.krs_done = self.nazwa_done = self.regon_done = self.nip_done = self.data_done = False

        self.regon_full = self.nip_full = ""

    def extract_pdf(self):
        subprocess.call(["C:\\WINDOWS\\system32\\WindowsPowerShell\\v1.0\\powershell.exe", ";",
                         "venv\\Scripts\\python.exe",
                         "venv\\Scripts\\pdf2txt.py",
                         f"'PDF\\{self.pdf_file}' -o 'txt\\{self.pdf_file}.txt'"])
        print(f"\n>>> Extracting text from '{self.pdf_file}'")

    def get_date(self):
        data = datetime.datetime.now()
        return f"{data.day}.{data.month}.{data.year}"

    def open_txt(self):  # returns string object from .txt
        with open(f"txt\\{self.txt_file}", "r", encoding="utf-8") as file:
            self.txt_string = [line.rstrip('\n') for line in file]

    def get_txt_string(self):  # debug only
        for counter, line in enumerate(self.txt_string):
            print(counter, line)
        return self.txt_string

    def parse_txt(self):
        """Loops through the txt file and maps distinct values to the class variables"""
        lines = self.txt_string
        for line in lines:
            if line == "Oznaczenie sądu":
                index = lines.index(line) + 4
                oznaczenie = lines[index]
                check = lines[index + 1]
                if check is not "":  # checks for a newline
                    oznaczenie += f" {check}"
                self.oznaczenie = oznaczenie
            if line.startswith("kraj "):
                index = lines.index(line)
                try:
                    kraj, woj, powiat, gmina, miejsc = lines[index].split(",")
                    self.woj = woj.strip("woj. ")
                    self.powiat = powiat.strip("  powiat ")
                    self.gmina = gmina.strip("  gmina ")
                    miejsc = miejsc.strip("  miejsc.")
                    self.miejsc = miejsc
                    if miejsc == "":  # checks for a newline
                        index = index + 1
                        self.miejsc = lines[index]
                except ValueError as error:
                    print(f">ERROR: {error}")
            if line.startswith("Numer KRS"):
                krs = line.split()[-1]
                if len(krs) is not 10:
                    print(">!!! INVALID KRS !!!\n")
                    krs = "ERROR"
                self.krs = krs
            if line.startswith("3.") and not self.nazwa_parsed:
                self.nazwa_parsed = True
                index = lines.index(line) + 2
                self.nazwa = lines[index]
            if line.startswith("2.") and not self.reg_nip_parsed:
                self.reg_nip_parsed = True
                index = lines.index(line) + 2
                pola = lines[index]
                regon, nip = pola.split(",")
                regon = regon.strip("REGON: ")
                self.regon = regon.split()[0]
                self.regon_full = regon
                nip = nip.strip("  NIP: ")
                self.nip = nip.split()[0].strip("---")
                self.nip_full = nip

    def get_parsed_data(self):  # debug only
        return self.krs, self.oznaczenie, self.woj, self.powiat, self.gmina,\
               self.miejsc, self.nazwa, self.regon, self.nip

    def get_formatted_data(self):
        print(
            f">>> Parsing... '{self.filename}'\n"
            f"> Nazwa sądu: {self.oznaczenie},\n> Województwo: {self.woj}, Powiat: {self.powiat},"
            f" Gmina: {self.gmina}, Miejscowość: {self.miejsc},"
            f"\n> Firma spółki: {self.nazwa},\n> Numer KRS: {self.krs}, REGON: {self.regon_full}, NIP: {self.nip_full}"
        )

    def parse_docx(self):
        """loops through the list of all cells in the document.
         Looks for a distinct element and inserts the new value in correct paragraph"""
        for cell in self.document.tables[0]._cells:
            if cell.paragraphs[0].text.startswith("1.") and not self.oznaczenie_done:
                cell.paragraphs[1].style.font.size = Pt(10)
                cell.paragraphs[1].text = self.oznaczenie
                self.oznaczenie_done = True

            if cell.paragraphs[0].text.startswith("2.") and not self.woj_done:
                cell.paragraphs[1].style.font.size = Pt(10)
                cell.paragraphs[1].text = self.woj
                self.woj_done = True

            if cell.paragraphs[0].text.startswith("3.") and not self.powiat_done:
                cell.paragraphs[1].style.font.size = Pt(10)
                cell.paragraphs[1].text = self.powiat
                self.powiat_done = True

            if cell.paragraphs[0].text.startswith("4.") and not self.gmina_done:
                cell.paragraphs[1].style.font.size = Pt(10)
                cell.paragraphs[1].text = self.gmina
                self.gmina_done = True

            if cell.paragraphs[0].text.startswith("5.") and not self.miejsc_done:
                cell.paragraphs[1].style.font.size = Pt(10)
                cell.paragraphs[1].text = self.miejsc
                self.miejsc_done = True

            if cell.paragraphs[0].text.startswith("<KRS>"):
                cell.paragraphs[0].text = ""
                if len(self.krs) is not 10:
                    print(">\n>!!! INVALID KRS !!!\n>")
                for liczba in self.krs:
                    paragraph = cell.paragraphs[0].add_run(liczba)
                    paragraph.font.size = Pt(10)
                    if int(liczba) < 5:
                        for i in range(15):
                            paragraph_space = cell.paragraphs[0].add_run(" ")
                            paragraph_space.font.size = Pt(2)
                    else:
                        for i in range(16):
                            paragraph_space = cell.paragraphs[0].add_run(" ")
                            paragraph_space.font.size = Pt(2)

            if cell.paragraphs[0].text.startswith("8.") and not self.nazwa_done:
                cell.paragraphs[1].style.font.size = Pt(10)
                cell.paragraphs[1].text = f"   {self.nazwa}"
                self.nazwa_done = True

            if cell.paragraphs[0].text.startswith("<NIP>") and not self.nip_done:
                cell.paragraphs[0].text = ""
                if len(self.nip) is not 10:
                    print(">\n>!!! INVALID NIP !!!\n>")
                for liczba in self.nip:
                    try:
                        paragraph = cell.paragraphs[0].add_run(liczba)
                        paragraph.font.size = Pt(10)
                        if int(liczba) < 5:  # takes care of number formatting
                            for i in range(15):
                                paragraph_space = cell.paragraphs[0].add_run(" ")
                                paragraph_space.font.size = Pt(2)
                        else:
                            for i in range(15):
                                paragraph_space = cell.paragraphs[0].add_run(" ")
                                paragraph_space.font.size = Pt(2)
                    except ValueError:
                        cell.paragraphs[0].add_run("")
                self.nip_done = True

            if cell.paragraphs[0].text.startswith("<REGON>") and not self.regon_done:
                cell.paragraphs[0].text = ""
                if len(self.regon) is not 9:
                    print(">\n>!!! INVALID REGON !!!\n>")
                for liczba in self.regon:
                    try:
                        paragraph = cell.paragraphs[0].add_run(liczba)
                        paragraph.font.size = Pt(10)
                        if int(liczba) < 5:
                            for i in range(15):
                                paragraph_space = cell.paragraphs[0].add_run(" ")
                                paragraph_space.font.size = Pt(2)
                        else:
                            for i in range(16):
                                paragraph_space = cell.paragraphs[0].add_run(" ")
                                paragraph_space.font.size = Pt(2)
                    except ValueError:
                        cell.paragraphs[0].add_run("")
                self.regon_done = True

        for cell in self.document.tables[2]._cells:
            if cell.paragraphs[0].text.startswith("DZ.MI"):  # and not self.data_done:
                cell.paragraphs[0].text = self.date

    def save_docx(self):
        os.chdir("DOCX\\")
        self.document.save(f"{self.filename}.docx")
        os.chdir("..\\")
        print(f">>> File '{self.filename}.docx' created successfully! <\n")

    def clear_temp(self):  # removes .txt file
        os.remove(f"txt\\{self.txt_file}")

    def parse_all(self):
        #  self.extract_pdf()  # extracts unicode text from pdf file
        self.open_txt()  # initiates temp text file
        self.parse_txt()  # parses temp text file
        self.get_formatted_data()  # console info
        self.parse_docx()  # inserts data into docx document
        self.save_docx()  # saves the .docx document


if __name__ == "__main__":  # debug
    parser = DocParser("..\\example.pdf")
    parser.extract_pdf()
    parser.open_txt()
    os.remove("example.pdf.txt")
    parser.parse_txt()
    parser.get_formatted_data()
    parser.parse_docx()
