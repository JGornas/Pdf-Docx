from pdf_parser import PdfParser
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os


class DocxParser:
    def __init__(self, pdf_datafields, docx_filename="template.docx"):
        self.document = Document(docx_filename)

        self.datafields = pdf_datafields  # Pdf_parser returned dictionary.
        self.cells = ["Oznaczenie sądu", "Siedziba_Wojewodztwo", "Siedziba_Powiat", "Siedziba_Gmina",
                      "Siedziba_Miejscowosc", "KRS", "Firma, pod którą spółka działa", "NIP", "REGON", "Current_date"]
        self.locked_cells = {key: False for key in self.cells}  # Locks document cell after inserting data.
        self.date = datetime.now()

    def parse_cell(self, cell, cell_name, startswith):
        def parse_default():
            cell.paragraphs[1].style.font.size = Pt(10)
            cell.paragraphs[1].text = self.datafields[cell_name]

        def parse_krs():
            cell.paragraphs[0].text = ""
            if len(self.datafields[cell_name]) is not 10:
                print("> No KRS number.")
            for liczba in self.datafields[cell_name]:
                paragraph = cell.paragraphs[0].add_run(liczba)
                paragraph.font.size = Pt(10)
                if int(liczba) < 5:
                    for i in range(12):
                        paragraph_space = cell.paragraphs[0].add_run(" ")
                        paragraph_space.font.size = Pt(3)
                else:
                    for i in range(13):
                        paragraph_space = cell.paragraphs[0].add_run(" ")
                        paragraph_space.font.size = Pt(3)

        def parse_nazwa():
            cell.paragraphs[1].style.font.size = Pt(10)
            cell.paragraphs[1].text = f"   {self.datafields[cell_name]}"

        def parse_nip():
            cell.paragraphs[0].text = ""
            if len(self.datafields[cell_name]) is not 10:
                print("> No NIP.")
                self.datafields[cell_name] = "               "
            for liczba in self.datafields[cell_name]:
                try:
                    paragraph = cell.paragraphs[0].add_run(liczba)
                    paragraph.font.size = Pt(10)
                    for i in range(12):
                        paragraph_space = cell.paragraphs[0].add_run(" ")
                        paragraph_space.font.size = Pt(3)
                except ValueError:
                    cell.paragraphs[0].add_run("")

        def parse_regon():
            cell.paragraphs[0].text = ""
            if len(self.datafields[cell_name]) is not 9:
                print("> No REGON.")
            for liczba in self.datafields[cell_name]:
                try:
                    paragraph = cell.paragraphs[0].add_run(liczba)
                    paragraph.font.size = Pt(10)
                    if int(liczba) < 5:  # this block controls the number of whitespaces between numbers
                        for i in range(12):
                            paragraph_space = cell.paragraphs[0].add_run(" ")
                            paragraph_space.font.size = Pt(3)
                    else:
                        for i in range(13):
                            paragraph_space = cell.paragraphs[0].add_run(" ")
                            paragraph_space.font.size = Pt(3)
                except ValueError:
                    cell.paragraphs[0].add_run("")

        def parse_date():
            cell.paragraphs[0].text = f"{self.date.day}.{self.date.month}.{self.date.year}"

        start = {"1.": parse_default, "2.": parse_default, "3.": parse_default,
                 "4.": parse_default, "5.": parse_default, "<KRS>": parse_krs,
                 "8.": parse_nazwa, "<NIP>": parse_nip, "<REGON>": parse_regon,
                 "DZ.MI": parse_date}

        if cell.paragraphs[0].text.startswith(startswith):
            start[startswith]()
            self.locked_cells[cell_name] = True

    def parse_document(self):
        startswith_name = [("1.", "Oznaczenie sądu"), ("2.", "Siedziba_Wojewodztwo"), ("3.", "Siedziba_Powiat"),
                           ("4.", "Siedziba_Gmina"), ("5.", "Siedziba_Miejscowosc"),
                           ("<KRS>", "KRS"), ("8.", "Firma, pod którą spółka działa"),
                           ("<NIP>", "NIP"), ("<REGON>", "REGON"), ("DZ.MI", "Current_date")]

        for table in range(len(self.document.tables)):
            try:
                for cell in self.document.tables[table]._cells:
                    for start_name in startswith_name:
                        if not self.locked_cells[start_name[1]]:
                            self.parse_cell(cell, start_name[1], start_name[0])
            except IndexError:
                pass

    def save_document(self, path="test.docx"):
        os.makedirs("docx", exist_ok=True)
        self.document.save(os.path.join("docx", path))


if __name__ == "__main__":
    parser = PdfParser()
    parser.load_pdf(env="venv")
    datafields = parser.parse_paragraphs()
    docx = DocxParser(datafields)
    docx.parse_document()
    docx.save_document()
